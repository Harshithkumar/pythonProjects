from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Create a new document
doc = Document()

# Title styling similar to the image
doc_title = doc.add_heading(level=0)
doc_title_run = doc_title.add_run("Harshith Kumar G S")
doc_title_run.font.size = Pt(24)
doc_title_run.font.bold = True
doc_title_run.font.color.rgb = RGBColor(54, 54, 150)
doc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Subtitle or tagline
doc_subtitle = doc.add_paragraph()
doc_subtitle_run = doc_subtitle.add_run("Dynamic and Passionate Generalist Product Manager")
doc_subtitle_run.font.size = Pt(12)
doc_subtitle_run.font.color.rgb = RGBColor(54, 54, 150)
doc_subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Contact Details Section
contact_section = doc.add_paragraph()
contact_section.add_run("Email: ").bold = True
contact_section.add_run("harsh.harshith@gmail.com  |  ")
contact_section.add_run("Phone: ").bold = True
contact_section.add_run("+91-9731947744  |  ")
contact_section.add_run("LinkedIn: ").bold = True
contact_section.add_run("linkedin.com/in/harshith-kumar-gs-11b64753/")
contact_section.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Line break
doc.add_paragraph("")

# Main Sections
sections = [("EXPERIENCE", [
                "Staff/Lead Engineer in Test - XPERI (2018 – Present): Spearheaded QA automation and collaborated with product owners for product roadmap planning.",
                "Senior Engineer and Team Lead - Rovi Corp (2016 – 2018): Led teams on TotalGuide XD and FanTV projects, increasing user engagement by 30%.",
                "Software Engineer - July Systems (CISCO) (2012 – 2016): Delivered seamless app integration for global brands like ESPN."
            ]),
            ("EDUCATION", [
                "M.Tech, BITS Pilani",
                "B.E., Computer Science"
            ]),
            ("CERTIFICATIONS", [
                "Certified Scrum Product Owner (CSPO)",
                "Certified Scrum Master (CSM)",
                "Product Management Training, Rethink Systems"
            ]),
            ("SKILLS", [
                "Product Management, Roadmap planning, API integration",
                "Agile Methodologies (Scrum, Kanban), JIRA, Confluence"
            ])]

# Add sections to the document with styling
for section_title, section_content in sections:
    # Section Title
    section_heading = doc.add_heading(level=1)
    section_heading_run = section_heading.add_run(section_title)
    section_heading_run.font.size = Pt(14)
    section_heading_run.font.bold = True
    section_heading_run.font.color.rgb = RGBColor(54, 54, 150)

    # Section Content
    for item in section_content:
        content_paragraph = doc.add_paragraph(style='List Bullet')
        content_run = content_paragraph.add_run(item)
        content_run.font.size = Pt(11)

# Save the document
styled_resume_path = "/mnt/data/Harshith_PM_Styled_Resume.docx"
doc.save(styled_resume_path)

styled_resume_path
